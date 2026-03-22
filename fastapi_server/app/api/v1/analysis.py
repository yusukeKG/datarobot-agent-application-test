# Copyright 2025 DataRobot, Inc.
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#   http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""Power-consumption analysis endpoint with 3-agent sequential workflow.

Agent 1 calls the CrewAI divergence-analysis agent via the DRUM endpoint.
Agents 2 and 3 still return mock responses (to be implemented later).
"""

from __future__ import annotations

import io
import json
import logging
import uuid as uuidpkg
from typing import Any, AsyncIterator

from datarobot.auth.session import AuthCtx
from datarobot.auth.typing import Metadata
from fastapi import APIRouter, Depends, HTTPException, Request
from fastapi.responses import StreamingResponse
from openai import AsyncOpenAI
from pydantic import BaseModel, Field

from app.analysis_duckdb import delete_timeseries, load_timeseries, save_timeseries
from app.analysis_reports import (
    AnalysisReportCreate,
    AnalysisReportPublic,
    AnalysisReportSummary,
)
from app.auth.ctx import must_get_auth_ctx
from app.deps import Deps
from app.users.user import User, UserRepository

logger = logging.getLogger(__name__)

analysis_router = APIRouter(prefix="/analysis", tags=["analysis"])


async def _get_current_user(user_repo: UserRepository, user_id: int) -> User:
    current_user = await user_repo.get_user(user_id=user_id)
    if not current_user:
        raise HTTPException(status_code=401, detail="User not found")
    return current_user


# ---------------------------------------------------------------------------
# Request / helper models
# ---------------------------------------------------------------------------


class AnomalyItem(BaseModel):
    timestamp: str
    power: float
    power_prediction: float
    diff: float
    diff_pct: float


class ChartDataRow(BaseModel):
    """A single row of chart data sent from the frontend."""

    timestamp: str = ""
    temperature: float | None = None
    fluidTemperature: float | None = None
    pressure: float | None = None
    power: float | None = None
    powerPrediction: float | None = None
    flow: float | None = None


class AnalysisRequest(BaseModel):
    """Payload sent by the sensors page."""

    start_date: str = Field(..., description="分析期間の開始日 (yyyy-MM-dd)")
    end_date: str = Field(..., description="分析期間の終了日 (yyyy-MM-dd)")
    anomaly_points: list[AnomalyItem] = Field(
        default_factory=list,
        description="予実乖離が閾値を超えた異常データポイント",
    )
    total_data_points: int = Field(0, description="期間内の全データポイント数")
    chart_data: list[ChartDataRow] = Field(
        default_factory=list,
        description="チャートに表示されている時系列データ全件",
    )


# ---------------------------------------------------------------------------
# SSE helper
# ---------------------------------------------------------------------------


def _sse_event(event: str, data: dict[str, Any]) -> str:
    """Format a single SSE frame."""
    return f"event: {event}\ndata: {json.dumps(data, ensure_ascii=False)}\n\n"


# ---------------------------------------------------------------------------
# Agent 1: Real LLM call via DRUM endpoint
# ---------------------------------------------------------------------------

_NO_ANOMALY_REPORT = (
    "## 予実乖離分析レポート\n\n"
    "**対象期間**: {start_date} ～ {end_date}\n\n"
    "この期間において、電力消費量の実績値と DataRobot 予測値の間に"
    "有意な乖離は検出されませんでした。\n\n"
    "設備は正常な稼働範囲内にあると判断されます。"
)


async def _agent_1_divergence_analysis(
    req: AnalysisRequest,
    agent_endpoint: str,
    api_token: str,
) -> str:
    """Agent 1: 予実乖離の時系列分析（LLM Gateway 経由）."""
    if not req.anomaly_points:
        return _NO_ANOMALY_REPORT.format(
            start_date=req.start_date, end_date=req.end_date
        )

    payload = {
        "start_date": req.start_date,
        "end_date": req.end_date,
        "total_data_points": req.total_data_points,
        "anomaly_points": [ap.model_dump() for ap in req.anomaly_points],
    }

    client = AsyncOpenAI(
        base_url=agent_endpoint,
        api_key=api_token,
        default_headers={"Authorization": f"Bearer {api_token}"},
    )
    try:
        response = await client.chat.completions.create(
            model="custom-model",
            messages=[
                {"role": "user", "content": json.dumps(payload, ensure_ascii=False)},
            ],
            stream=False,
        )
        return response.choices[0].message.content or ""
    except Exception:
        logger.exception("Agent 1 LLM call failed")
        raise


async def _agent_2_past_cases(
    analysis_summary: str,
    req: AnalysisRequest,
    agent_endpoint: str,
    api_token: str,
) -> str:
    """Agent 2: 過去の類似事例検索（LLM Gateway 経由）."""
    if not req.anomaly_points:
        return (
            "## 過去事例検索結果\n\n"
            "現在の分析では有意な異常が検出されていないため、"
            "過去の類似事例検索は省略されました。"
        )

    payload = {
        "agent_type": "past_case_search",
        "analysis_summary": analysis_summary,
        "anomaly_points": [ap.model_dump() for ap in req.anomaly_points],
    }

    client = AsyncOpenAI(
        base_url=agent_endpoint,
        api_key=api_token,
        default_headers={"Authorization": f"Bearer {api_token}"},
    )
    try:
        response = await client.chat.completions.create(
            model="custom-model",
            messages=[
                {"role": "user", "content": json.dumps(payload, ensure_ascii=False)},
            ],
            stream=False,
        )
        return response.choices[0].message.content or ""
    except Exception:
        logger.exception("Agent 2 LLM call failed")
        raise


async def _agent_3_maintenance_actions(
    analysis_summary: str,
    past_cases: str,
    req: AnalysisRequest,
    agent_endpoint: str,
    api_token: str,
) -> str:
    """Agent 3: 保守アクション提案（LLM Gateway 経由）."""
    if len(req.anomaly_points) == 0:
        return (
            "## 推奨アクション\n\n"
            "現時点で緊急の対応は不要です。\n\n"
            "定期点検スケジュールに従い、次回点検時に"
            "センサーの校正確認を行うことを推奨します。"
        )

    payload = {
        "agent_type": "maintenance_action",
        "analysis_summary": analysis_summary,
        "past_cases": past_cases,
    }

    client = AsyncOpenAI(
        base_url=agent_endpoint,
        api_key=api_token,
        default_headers={"Authorization": f"Bearer {api_token}"},
    )
    try:
        response = await client.chat.completions.create(
            model="custom-model",
            messages=[
                {"role": "user", "content": json.dumps(payload, ensure_ascii=False)},
            ],
            stream=False,
        )
        return response.choices[0].message.content or ""
    except Exception:
        logger.exception("Agent 3 LLM call failed")
        raise


# ---------------------------------------------------------------------------
# SSE streaming endpoint
# ---------------------------------------------------------------------------


async def _run_analysis_pipeline(
    req: AnalysisRequest,
    agent_endpoint: str,
    api_token: str,
    deps: Deps,
    user_uuid: uuidpkg.UUID | None = None,
) -> AsyncIterator[str]:
    """Run the 3-agent pipeline, yielding SSE events."""

    agents = [
        {
            "agent_id": 1,
            "title": "予実乖離の時系列分析",
            "description": "予測と実績の乖離パターンを時間軸方向に分析しています…",
        },
        {
            "agent_id": 2,
            "title": "過去事例の検索",
            "description": "過去の保守報告書から類似事例を検索しています…",
        },
        {
            "agent_id": 3,
            "title": "保守アクションの提案",
            "description": "分析結果と過去事例に基づき推奨アクションを生成しています…",
        },
    ]

    # Emit all agents as "pending" first
    for agent in agents:
        yield _sse_event(
            "agent_step",
            {
                "agent_id": agent["agent_id"],
                "title": agent["title"],
                "description": agent["description"],
                "status": "pending",
                "content": "",
            },
        )

    # --- Agent 1 ---
    yield _sse_event(
        "agent_step",
        {
            "agent_id": 1,
            "title": agents[0]["title"],
            "status": "running",
            "content": "",
        },
    )
    try:
        agent1_result = await _agent_1_divergence_analysis(
            req, agent_endpoint, api_token
        )
    except Exception as exc:
        agent1_result = f"## エラー\n\nAgent 1 の実行中にエラーが発生しました: {exc}"
    yield _sse_event(
        "agent_step",
        {
            "agent_id": 1,
            "title": agents[0]["title"],
            "status": "completed",
            "content": agent1_result,
        },
    )

    # --- Agent 2 ---
    yield _sse_event(
        "agent_step",
        {
            "agent_id": 2,
            "title": agents[1]["title"],
            "status": "running",
            "content": "",
        },
    )
    try:
        agent2_result = await _agent_2_past_cases(
            agent1_result, req, agent_endpoint, api_token
        )
    except Exception as exc:
        agent2_result = f"## エラー\n\nAgent 2 の実行中にエラーが発生しました: {exc}"
    yield _sse_event(
        "agent_step",
        {
            "agent_id": 2,
            "title": agents[1]["title"],
            "status": "completed",
            "content": agent2_result,
        },
    )

    # --- Agent 3 ---
    yield _sse_event(
        "agent_step",
        {
            "agent_id": 3,
            "title": agents[2]["title"],
            "status": "running",
            "content": "",
        },
    )
    try:
        agent3_result = await _agent_3_maintenance_actions(
            agent1_result, agent2_result, req, agent_endpoint, api_token
        )
    except Exception as exc:
        agent3_result = f"## エラー\n\nAgent 3 の実行中にエラーが発生しました: {exc}"
    yield _sse_event(
        "agent_step",
        {
            "agent_id": 3,
            "title": agents[2]["title"],
            "status": "completed",
            "content": agent3_result,
        },
    )

    # --- Save results to SQLite + DuckDB ---
    report_uuid: str | None = None
    try:
        # Determine anomaly timestamps for marking
        anomaly_timestamps = {ap.timestamp for ap in req.anomaly_points}

        # Convert chart_data rows for DuckDB
        ts_rows = [
            {
                "timestamp": r.timestamp,
                "temperature": r.temperature,
                "fluid_temperature": r.fluidTemperature,
                "pressure": r.pressure,
                "power": r.power,
                "power_prediction": r.powerPrediction,
                "flow": r.flow,
                "is_anomaly": r.timestamp in anomaly_timestamps,
            }
            for r in req.chart_data
        ]

        new_uuid = uuidpkg.uuid4()
        duckdb_table = save_timeseries(new_uuid, ts_rows)

        report_data = AnalysisReportCreate(
            uuid=new_uuid,
            user_uuid=user_uuid,
            start_date=req.start_date,
            end_date=req.end_date,
            total_data_points=req.total_data_points,
            anomaly_count=len(req.anomaly_points),
            divergence_report=agent1_result,
            past_cases_report=agent2_result,
            maintenance_actions_report=agent3_result,
            duckdb_table_name=duckdb_table,
        )
        saved = await deps.analysis_report_repo.create(report_data)
        report_uuid = str(saved.uuid)
    except Exception:
        logger.exception("Failed to save analysis report")

    # --- Done ---
    yield _sse_event(
        "analysis_complete",
        {"status": "done", "report_uuid": report_uuid},
    )


@analysis_router.post("/power-consumption")
async def analyze_power_consumption(
    req: AnalysisRequest,
    request: Request,
    auth_ctx: AuthCtx[Metadata] = Depends(must_get_auth_ctx),
) -> StreamingResponse:
    """Run the 3-agent power-consumption analysis pipeline.

    Returns an SSE stream with ``agent_step`` events for each agent
    and an ``analysis_complete`` event when the pipeline finishes.
    """
    deps: Deps = request.app.state.deps
    current_user = await _get_current_user(
        deps.user_repo, int(auth_ctx.user.id)
    )
    return StreamingResponse(
        _run_analysis_pipeline(
            req,
            agent_endpoint=deps.config.agent_endpoint,
            api_token=deps.config.datarobot_api_token,
            deps=deps,
            user_uuid=current_user.uuid,
        ),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


# ---------------------------------------------------------------------------
# CRUD endpoints for analysis reports
# ---------------------------------------------------------------------------


@analysis_router.get("/reports")
async def list_reports(
    request: Request,
    auth_ctx: AuthCtx[Metadata] = Depends(must_get_auth_ctx),
) -> list[AnalysisReportSummary]:
    """Return all analysis reports for the current user."""
    deps: Deps = request.app.state.deps
    current_user = await _get_current_user(
        deps.user_repo, int(auth_ctx.user.id)
    )
    reports = await deps.analysis_report_repo.list_by_user(current_user.uuid)
    return [
        AnalysisReportSummary(
            uuid=r.uuid,
            created_at=r.created_at,
            start_date=r.start_date,
            end_date=r.end_date,
            total_data_points=r.total_data_points,
            anomaly_count=r.anomaly_count,
        )
        for r in reports
    ]


@analysis_router.get("/reports/{report_uuid}")
async def get_report(
    report_uuid: uuidpkg.UUID,
    request: Request,
    auth_ctx: AuthCtx[Metadata] = Depends(must_get_auth_ctx),
) -> AnalysisReportPublic:
    """Return full details for a single analysis report."""
    deps: Deps = request.app.state.deps
    current_user = await _get_current_user(
        deps.user_repo, int(auth_ctx.user.id)
    )
    report = await deps.analysis_report_repo.get_by_uuid(report_uuid)
    if not report or report.user_uuid != current_user.uuid:
        raise HTTPException(status_code=404, detail="Report not found")
    return AnalysisReportPublic(**report.model_dump())


@analysis_router.get("/reports/{report_uuid}/timeseries")
async def get_report_timeseries(
    report_uuid: uuidpkg.UUID,
    request: Request,
    auth_ctx: AuthCtx[Metadata] = Depends(must_get_auth_ctx),
) -> list[dict[str, Any]]:
    """Return the timeseries chart data stored in DuckDB."""
    deps: Deps = request.app.state.deps
    current_user = await _get_current_user(
        deps.user_repo, int(auth_ctx.user.id)
    )
    report = await deps.analysis_report_repo.get_by_uuid(report_uuid)
    if not report or report.user_uuid != current_user.uuid:
        raise HTTPException(status_code=404, detail="Report not found")
    return load_timeseries(report.uuid)


@analysis_router.delete("/reports/{report_uuid}")
async def delete_report(
    report_uuid: uuidpkg.UUID,
    request: Request,
    auth_ctx: AuthCtx[Metadata] = Depends(must_get_auth_ctx),
) -> dict[str, str]:
    """Delete an analysis report (SQLite + DuckDB)."""
    deps: Deps = request.app.state.deps
    current_user = await _get_current_user(
        deps.user_repo, int(auth_ctx.user.id)
    )
    report = await deps.analysis_report_repo.get_by_uuid(report_uuid)
    if not report or report.user_uuid != current_user.uuid:
        raise HTTPException(status_code=404, detail="Report not found")
    delete_timeseries(report.uuid)
    await deps.analysis_report_repo.delete(report_uuid)
    return {"status": "deleted"}


# ---------------------------------------------------------------------------
# Word document download
# ---------------------------------------------------------------------------


def _generate_word_report(
    report: AnalysisReportPublic,
    ts_rows: list[dict[str, Any]],
) -> bytes:
    """Build a .docx report with a matplotlib chart and agent outputs."""
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.add_heading("予実乖離分析レポート", level=0)

    # Basic info table
    table = doc.add_table(rows=4, cols=2)
    table.style = "Light Grid Accent 1"
    cells = table.rows[0].cells
    cells[0].text = "対象期間"
    cells[1].text = f"{report.start_date} ～ {report.end_date}"
    cells = table.rows[1].cells
    cells[0].text = "データポイント数"
    cells[1].text = str(report.total_data_points)
    cells = table.rows[2].cells
    cells[0].text = "異常ポイント数"
    cells[1].text = str(report.anomaly_count)
    cells = table.rows[3].cells
    cells[0].text = "分析日時"
    cells[1].text = report.created_at.strftime("%Y-%m-%d %H:%M")

    doc.add_paragraph("")

    # Chart
    if ts_rows:
        fig, ax = plt.subplots(figsize=(10, 4))
        timestamps = [r["timestamp"] for r in ts_rows]
        power = [r.get("power") for r in ts_rows]
        prediction = [r.get("power_prediction") for r in ts_rows]

        ax.plot(range(len(timestamps)), power, color="#f97316", label="実績", linewidth=1)
        ax.plot(
            range(len(timestamps)),
            prediction,
            color="#3b82f6",
            label="予測",
            linewidth=1,
        )

        # Anomaly markers
        anomaly_indices = [
            i for i, r in enumerate(ts_rows) if r.get("is_anomaly")
        ]
        anomaly_power = [power[i] for i in anomaly_indices]
        if anomaly_indices:
            ax.scatter(anomaly_indices, anomaly_power, color="red", s=30, zorder=5, label="異常")

        ax.set_ylabel("電力消費量 (kWh)")
        ax.legend(fontsize=8)

        # Show sparse x-tick labels
        n_labels = min(10, len(timestamps))
        step = max(1, len(timestamps) // n_labels)
        ax.set_xticks(range(0, len(timestamps), step))
        ax.set_xticklabels(
            [timestamps[i] for i in range(0, len(timestamps), step)],
            rotation=45,
            ha="right",
            fontsize=7,
        )
        fig.tight_layout()

        img_buf = io.BytesIO()
        fig.savefig(img_buf, format="png", dpi=150)
        plt.close(fig)
        img_buf.seek(0)

        doc.add_heading("電力消費量チャート", level=2)
        doc.add_picture(img_buf, width=Inches(6))
        doc.add_paragraph("")

    # Agent reports
    sections = [
        ("Agent 1: 予実乖離の時系列分析", report.divergence_report),
        ("Agent 2: 過去事例の検索", report.past_cases_report),
        ("Agent 3: 保守アクションの提案", report.maintenance_actions_report),
    ]
    for title, content in sections:
        doc.add_heading(title, level=2)
        for line in content.split("\n"):
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph("")
            elif stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=4)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=3)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=2)
            else:
                doc.add_paragraph(stripped)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@analysis_router.get("/reports/{report_uuid}/download")
async def download_report(
    report_uuid: uuidpkg.UUID,
    request: Request,
    auth_ctx: AuthCtx[Metadata] = Depends(must_get_auth_ctx),
) -> StreamingResponse:
    """Download the analysis report as a Word (.docx) file."""
    deps: Deps = request.app.state.deps
    current_user = await _get_current_user(
        deps.user_repo, int(auth_ctx.user.id)
    )
    report = await deps.analysis_report_repo.get_by_uuid(report_uuid)
    if not report or report.user_uuid != current_user.uuid:
        raise HTTPException(status_code=404, detail="Report not found")

    ts_rows = load_timeseries(report_uuid)
    report_public = AnalysisReportPublic(**report.model_dump())
    docx_bytes = _generate_word_report(report_public, ts_rows)

    filename = f"analysis_report_{report.start_date}_{report.end_date}.docx"
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
